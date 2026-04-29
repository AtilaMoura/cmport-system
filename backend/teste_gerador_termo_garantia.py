"""
Teste isolado do gerador de Termo de Garantia.
Roda sem banco — usa dados mockados.
Gera termo_teste.docx e (se LibreOffice disponível) termo_teste.pdf

Uso:
  cd backend
  venv\\Scripts\\activate
  python teste_gerador_termo_garantia.py
"""
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "app", "assets", "termo_garantia_template.docx")
OUT_DOCX = os.path.join(os.path.dirname(__file__), "termo_teste.docx")
OUT_PDF  = os.path.join(os.path.dirname(__file__), "termo_teste.pdf")

# ─── Dados mockados ────────────────────────────────────────────────────────────
EMPRESA_NOME     = "CMPORT Sistemas Eletrônicos de Segurança"
CONDO_NOME       = "Condomínio Edifício São Bento Green Park"
ENDERECO_STR     = "R. Tupiguaes, 123, Vila Bianca, São Paulo"
TIPO_SERVICO     = "manutencao"   # "manutencao" → -M  |  "assistencia" → -A
NUMERO_NOTA_RAW  = "7710"
NUMERO_OS        = "71036107"
PRAZO_MESES      = 12
DATA_SERVICO     = date(2026, 3, 18)
DATA_INICIO      = date(2026, 3, 18)
DATA_FIM         = date(2027, 3, 18)
HOJE             = date(2026, 4, 29)

# Lista de produtos do checklist (como viria do frontend)
PRODUTOS = [
    "2x Câmera HiKVision HWC-B120H",
    "1x NVDS 4 Canais",
    "3x Sensor de Presença PPA",
    "1x Central de Alarme Intelbras",
]

# ─── Helpers ───────────────────────────────────────────────────────────────────
_EXTENSO = {3: "três", 6: "seis", 12: "doze", 24: "vinte e quatro"}

def fmt_date(d):
    return d.strftime('%d/%m/%Y')

def fmt_data_extenso(d):
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"

def fmt_numero_nota(numero_raw, tipo_servico):
    clean = str(numero_raw).strip().replace('.', '')
    m = re.match(r'^(\d+)', clean)
    if not m:
        return numero_raw
    num = int(m.group(1))
    digits = f"{num:09d}"
    formatted = f"{digits[0:3]}.{digits[3:6]}.{digits[6:9]}"
    sufixo = '-A' if str(tipo_servico).lower() == 'assistencia' else '-M'
    return f"{formatted}{sufixo}"

def para_text(para):
    return ''.join(r.text for r in para.runs)

def merge_all(para, new_text):
    if not para.runs:
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''

def set_cliente_endereco(para, nome, endereco):
    runs = para.runs
    br_idx = next((i for i, r in enumerate(runs)
                   if r._r.find(qn('w:br')) is not None), -1)
    if br_idx == -1:
        merge_all(para, f"Cliente: {nome}. Endereço: {endereco}.")
        return
    runs[0].text = f"Cliente: {nome}."
    for r in runs[1:br_idx]:
        r.text = ''
    after = runs[br_idx + 1:]
    if after:
        after[0].text = f"Endereço: {endereco}."
        for r in after[1:]:
            r.text = ''

def set_normal_runs(para, new_value):
    normal_runs = [r for r in para.runs if r.bold is not True]
    if not normal_runs:
        return
    normal_runs[0].text = f' {new_value}'
    for r in normal_runs[1:]:
        r.text = ''

def set_prazo(para, prazo_fmt, data_inicio, data_fim):
    runs = para.runs
    in_label = True
    intro_passed = False
    bold_prazo = []
    normal_datas = None
    for run in runs:
        if in_label:
            if run.bold is True:
                continue
            else:
                in_label = False
                intro_passed = True
                continue
        elif intro_passed and run.bold is True:
            bold_prazo.append(run)
        elif intro_passed and run.bold is not True:
            normal_datas = run
    if bold_prazo:
        bold_prazo[0].text = f'{prazo_fmt} meses'
        for r in bold_prazo[1:]:
            r.text = ''
    if normal_datas:
        normal_datas.text = f', com início em {data_inicio} e término em {data_fim}.'

def remove_para(para):
    p = para._element
    p.getparent().remove(p)

def remover_quebras_pagina(doc):
    """Remove pageBreakBefore e runs com w:br type='page' de todos os parágrafos."""
    for para in doc.paragraphs:
        ppr = para._element.find(qn('w:pPr'))
        if ppr is not None:
            pb = ppr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                ppr.remove(pb)
        for run in para.runs:
            for br in run._r.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    run._r.remove(br)

# ─── Debug ─────────────────────────────────────────────────────────────────────
def debug_template(doc):
    print("\n" + "="*60)
    print("PARÁGRAFOS DO TEMPLATE (com texto)")
    print("="*60)
    for i, para in enumerate(doc.paragraphs):
        text = para_text(para)
        if not text.strip():
            continue
        print(f"\n[{i:2d}] TEXTO: {repr(text[:100])}")
        # Verifica pageBreakBefore
        ppr = para._element.find(qn('w:pPr'))
        if ppr is not None:
            pb = ppr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                val = pb.get(qn('w:val'), 'true')
                if val.lower() not in ('false', '0'):
                    print(f"      *** pageBreakBefore ATIVO ***")
            # keepWithNext
            kwn = ppr.find(qn('w:keepNext'))
            if kwn is not None:
                print(f"      keepWithNext: {kwn.get(qn('w:val'), 'true')}")
        for j, r in enumerate(para.runs):
            # Verifica page break dentro do run
            has_pb = any(br.get(qn('w:type')) == 'page'
                         for br in r._r.findall(qn('w:br')))
            pb_info = " *** PAGE BREAK NO RUN ***" if has_pb else ""
            print(f"      run[{j}] bold={str(r.bold):<5} size={r.font.size} "
                  f"text={repr(r.text[:50])}{pb_info}")
    print("="*60 + "\n")


# ─── Gerador principal ─────────────────────────────────────────────────────────
def gerar_docx():
    numero_nota_fmt  = fmt_numero_nota(NUMERO_NOTA_RAW, TIPO_SERVICO)
    prazo_extenso    = _EXTENSO.get(PRAZO_MESES, str(PRAZO_MESES))
    prazo_fmt        = f"{PRAZO_MESES:02d} ({prazo_extenso})"
    produto_desc     = " · ".join(PRODUTOS)
    hoje_str         = fmt_data_extenso(HOJE)
    data_servico_str = fmt_date(DATA_SERVICO)
    data_inicio_str  = fmt_date(DATA_INICIO)
    data_fim_str     = fmt_date(DATA_FIM)

    print(f"Nota fiscal formatada : {numero_nota_fmt}")
    print(f"Produto descrição     : {produto_desc}")
    print(f"Prazo                 : {prazo_fmt}")

    doc = Document(TEMPLATE_PATH)
    remover_quebras_pagina(doc)
    debug_template(doc)

    paras_remover = []
    for para in doc.paragraphs:
        text = para_text(para)
        if not text.strip():
            continue

        if "Paulo," in text and len(text) < 60:
            merge_all(para, f"São Paulo, {hoje_str}")
            print(f"  ✓ data: São Paulo, {hoje_str}")

        elif "Cliente:" in text and "Endere" in text:
            set_cliente_endereco(para, CONDO_NOME, ENDERECO_STR)
            print(f"  ✓ cliente/endereço")

        elif "consistente na" in text:
            aplicou = False
            for run in para.runs:
                if run.bold is True:
                    run.text = produto_desc
                    run.font.size = Pt(9)
                    aplicou = True
                    break
            if not aplicou:
                print(f"  ⚠ 'consistente na' encontrado mas nenhum run bold — runs: {[(r.bold, repr(r.text[:30])) for r in para.runs]}")
            else:
                print(f"  ✓ produto_desc aplicado com Pt(9)")

        elif "execu" in text and "servi" in text and "/" in text:
            set_normal_runs(para, data_servico_str)
            print(f"  ✓ data execução: {data_servico_str}")

        elif "Nota Fiscal:" in text:
            if numero_nota_fmt:
                set_normal_runs(para, f"nº {numero_nota_fmt}")
                print(f"  ✓ nota fiscal: {numero_nota_fmt}")
            else:
                paras_remover.append(para)

        elif "Ordem de Servi" in text:
            if NUMERO_OS:
                set_normal_runs(para, f"nº {NUMERO_OS}")
                print(f"  ✓ OS: {NUMERO_OS}")
            else:
                paras_remover.append(para)

        elif "Prazo de Garantia:" in text and "garantia concedida" in text:
            set_prazo(para, prazo_fmt, data_inicio_str, data_fim_str)
            print(f"  ✓ prazo: {prazo_fmt}")

        elif "André Moreira Rosa" in text:
            aplicou = False
            for run in para.runs:
                stripped = run.text.rstrip()
                if stripped and "André" not in stripped and "Diretor" not in stripped:
                    trailing = run.text[len(stripped):]
                    run.text = EMPRESA_NOME + trailing
                    aplicou = True
                    print(f"  ✓ assinatura empresa: '{run.text[:50]}'")
                    break
            if not aplicou:
                print(f"  ⚠ assinatura — nenhum run substituível. runs: {[(repr(r.text[:30])) for r in para.runs]}")

    for para in paras_remover:
        remove_para(para)

    return doc


if __name__ == "__main__":
    print(f"\nTemplate : {TEMPLATE_PATH}")
    print(f"Saída    : {OUT_DOCX}\n")

    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERRO: template não encontrado em {TEMPLATE_PATH}")
        exit(1)

    doc = gerar_docx()
    doc.save(OUT_DOCX)
    print(f"\n✓ DOCX salvo: {OUT_DOCX}")

    # Tenta converter para PDF via LibreOffice (disponível dentro do container ou se instalado localmente)
    lo = shutil.which("libreoffice") or shutil.which("soffice")
    if lo:
        tmp = tempfile.mkdtemp()
        try:
            result = subprocess.run(
                [lo, '--headless', '--convert-to', 'pdf', '--outdir', tmp, OUT_DOCX],
                check=True, capture_output=True, timeout=60
            )
            pdf_src = os.path.join(tmp, "termo_teste.pdf")
            shutil.copy(pdf_src, OUT_PDF)
            print(f"✓ PDF salvo : {OUT_PDF}")
        except subprocess.CalledProcessError as e:
            print(f"⚠ LibreOffice falhou: {e.stderr.decode()}")
        except Exception as e:
            print(f"⚠ Erro ao converter: {e}")
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    else:
        print("⚠ LibreOffice não encontrado no PATH.")
        print("  → Abra termo_teste.docx no Word para visualizar.")
        print("  → Ou rode dentro do container Docker onde LibreOffice está instalado.")
