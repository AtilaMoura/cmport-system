import io
import os
import shutil
import subprocess
import tempfile
from datetime import datetime
from sqlalchemy.orm import Session

from app.repositories.termo_garantia_repository import TermoGarantiaRepository
from app.models.configuracao_model import ConfiguracaoEmpresa

_EXTENSO = {3: "três", 6: "seis", 12: "doze", 24: "vinte e quatro"}
_EMPRESA_FALLBACK = "CMPORT Sistemas Eletrônicos de Segurança"
_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "termo_garantia_template.docx")


def _fmt_date(d) -> str:
    return d.strftime('%d/%m/%Y') if d else ""


def _fmt_numero_nota(numero_raw: str, tipo_servico) -> str:
    """
    Formata o número da nota para o padrão 000.000.094-A / 000.000.094-M.
    Extrai apenas os dígitos iniciais (ignora sufixos como '-2', '-A', etc.)
    e aplica sufixo baseado no tipo do serviço.
    """
    import re
    clean = str(numero_raw).strip().replace('.', '')
    m = re.match(r'^(\d+)', clean)
    if not m:
        return numero_raw
    num = int(m.group(1))
    digits = f"{num:09d}"
    formatted = f"{digits[0:3]}.{digits[3:6]}.{digits[6:9]}"
    valor_tipo = getattr(tipo_servico, 'value', str(tipo_servico)).lower()
    sufixo = '-A' if valor_tipo == 'assistencia' else '-M'
    return f"{formatted}{sufixo}"


def _fmt_data_extenso(d) -> str:
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _para_text(para) -> str:
    return ''.join(r.text for r in para.runs)


def _merge_all(para, new_text: str):
    """Substitui todo o texto no primeiro run, limpa os demais."""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''


def _set_cliente_endereco(para, nome: str, endereco: str):
    """
    P5 tem estrutura:
      runs 0-4: 'Cliente:', ' ', 'Condominio ...', 'NOME', '.'
      run  5:   '\n' — apenas <w:br/> (quebra de linha)
      runs 6-12: 'Endereço:', ' ', partes do endereço...
    Preserva o <w:br/> intacto entre cliente e endereço.
    Label em texto normal; valor (nome/endereço) em negrito.
    """
    from docx.oxml.ns import qn
    runs = para.runs
    br_idx = next((i for i, r in enumerate(runs) if r._r.find(qn('w:br')) is not None), -1)

    if br_idx == -1:
        _merge_all(para, f"Cliente: {nome}. Endereço: {endereco}.")
        return

    # Antes do break: "Cliente: " normal + nome bold
    before = runs[:br_idx]
    if len(before) >= 2:
        before[0].text = 'Cliente: '
        before[1].text = nome
        before[1].bold = True
        for r in before[2:]:
            r.text = ''
    elif len(before) == 1:
        before[0].text = f'Cliente: {nome}'
    # run[br_idx] não é tocado → mantém o <w:br/>

    # Após o break: "Endereço: " normal + endereço bold
    after = runs[br_idx + 1:]
    if len(after) >= 2:
        after[0].text = 'Endereço: '
        after[1].text = endereco
        after[1].bold = True
        for r in after[2:]:
            r.text = ''
    elif len(after) == 1:
        after[0].text = f'Endereço: {endereco}'


def _set_normal_runs(para, new_value: str):
    """Mantém runs bold intactos; substitui os runs não-bold pelo novo valor."""
    normal_runs = [r for r in para.runs if r.bold is not True]
    if not normal_runs:
        return
    normal_runs[0].text = f' {new_value}'
    for r in normal_runs[1:]:
        r.text = ''


def _set_prazo(para, prazo_fmt: str, data_inicio: str, data_fim: str):
    """
    P10 tem estrutura:
      run 0: B  'Prazo de Garantia:'          — label bold, não toca
      run 1: N  '\nA garantia concedida é de ' — intro c/ <w:br/>, não toca
      runs 2-6: B  '06', ' (', 'seis', ') ', 'meses' — bold prazo, mescla num só
      run 7: N  ', com início em ... e término em ...' — datas, atualiza
    """
    runs = para.runs
    in_label = True
    intro_passed = False
    bold_prazo = []
    normal_datas = None

    for run in runs:
        if in_label:
            if run.bold is True:
                continue  # ainda no label bold
            else:
                in_label = False
                intro_passed = True
                continue  # intro com <w:br/> — não modifica
        elif intro_passed and run.bold is True:
            bold_prazo.append(run)
        elif intro_passed and run.bold is not True:
            normal_datas = run

    if bold_prazo:
        bold_prazo[0].text = f'{prazo_fmt} meses'
        for r in bold_prazo[1:]:
            r.text = ''

    if normal_datas:
        normal_datas.text = f', com início em {data_inicio} e término em {data_fim}.'


def _remove_para(para):
    """Remove o parágrafo do documento."""
    p = para._element
    p.getparent().remove(p)


def _remover_quebras_pagina(doc):
    """Remove pageBreakBefore e runs com w:br type='page' de todos os parágrafos."""
    from docx.oxml.ns import qn
    for para in doc.paragraphs:
        ppr = para._element.find(qn('w:pPr'))
        if ppr is not None:
            pb = ppr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                ppr.remove(pb)
        for run in para.runs:
            for br in run._r.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    run._r.remove(br)


def _remover_paragrafos_vazios_iniciais(doc):
    """Remove os parágrafos vazios do início do body.
    No Word eles empurram o conteúdo para baixo do logo, mas no LibreOffice
    criam a faixa branca porque o top_margin já posiciona o corpo corretamente.
    """
    para_remover = []
    for para in doc.paragraphs:
        if not _para_text(para).strip():
            para_remover.append(para)
        else:
            break
    for para in para_remover:
        _remove_para(para)


def _ajustar_para_uma_pagina(doc):
    """Ajusta margens para caber em uma página sem sobrepor o logo do cabeçalho."""
    from docx.shared import Cm, Pt
    for section in doc.sections:
        # top_margin calibrado para começar logo abaixo do logo do header
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(1.5)
        section.footer_distance = Cm(0.75)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size > Pt(10):
                run.font.size = run.font.size - Pt(1)


class TermoGarantiaService:

    @staticmethod
    def gerar_pdf(db: Session, termo_id: int) -> io.BytesIO:
        from docx import Document
        from docx.shared import Pt

        termo = TermoGarantiaRepository.get_by_id(db, termo_id)
        if not termo:
            raise Exception("Termo de garantia não encontrado")

        servico = termo.servico
        condominio = servico.condominio

        empresa_obj = db.query(ConfiguracaoEmpresa).first()
        empresa_nome = (empresa_obj.nome if empresa_obj and empresa_obj.nome else _EMPRESA_FALLBACK)

        end = getattr(condominio, 'endereco', None)
        if end:
            partes = [p for p in [end.rua, end.numero, end.bairro, end.cidade] if p]
            endereco_str = ", ".join(partes)
        else:
            endereco_str = ""

        numero_nota = ""
        numero_nota_fmt = ""
        if getattr(servico, 'nota_fiscal', None):
            numero_nota = servico.nota_fiscal.numero_nota or ""
            if numero_nota:
                numero_nota_fmt = _fmt_numero_nota(numero_nota, servico.tipo)

        numero_os = servico.numero_os or ""
        prazo_extenso = _EXTENSO.get(termo.prazo_meses, str(termo.prazo_meses))
        prazo_fmt = f"{termo.prazo_meses:02d} ({prazo_extenso})"

        # Descrição do produto: usa produto_descricao (já formatado pelo checklist do frontend).
        # Fallback para report da OS se campo estiver vazio.
        produto_desc = termo.produto_descricao or ""
        if not produto_desc:
            if getattr(servico, 'ordem_servico', None):
                produto_desc = servico.ordem_servico.report or ""

        hoje_str = _fmt_data_extenso(datetime.now())
        data_servico_str = _fmt_date(servico.data_servico)
        data_inicio_str = _fmt_date(termo.data_inicio)
        data_fim_str = _fmt_date(termo.data_fim)

        doc = Document(_TEMPLATE_PATH)
        _remover_quebras_pagina(doc)
        _remover_paragrafos_vazios_iniciais(doc)
        _ajustar_para_uma_pagina(doc)

        paras_remover = []
        for para in doc.paragraphs:
            text = _para_text(para)
            if not text.strip():
                continue

            if "Paulo," in text and len(text) < 60:
                # "São Paulo, 27 de abril de 2026"
                _merge_all(para, f"São Paulo, {hoje_str}")

            elif "Cliente:" in text and "Endere" in text:
                # "Cliente: NOME\nEndereço: ENDEREÇO"
                _set_cliente_endereco(para, condominio.nome, endereco_str)

            elif "consistente na" in text:
                for run in para.runs:
                    if run.bold is True:
                        run.text = produto_desc
                        run.font.size = Pt(9)
                        break

            elif "execu" in text and "servi" in text and "/" in text:
                # "Data da execução do serviço: DD/MM/AAAA"
                _set_normal_runs(para, data_servico_str)

            elif "Nota Fiscal:" in text:
                if numero_nota_fmt:
                    _set_normal_runs(para, f"nº {numero_nota_fmt}")
                else:
                    paras_remover.append(para)

            elif "Ordem de Servi" in text:
                if numero_os:
                    _set_normal_runs(para, f"nº {numero_os}")
                else:
                    paras_remover.append(para)

            elif "Prazo de Garantia:" in text and "garantia concedida" in text:
                _set_prazo(para, prazo_fmt, data_inicio_str, data_fim_str)

            elif "André Moreira Rosa" in text:
                # Assinatura: substitui empresa preservando espaços de alinhamento
                for run in para.runs:
                    stripped = run.text.rstrip()
                    if stripped and "André" not in stripped and "Diretor" not in stripped:
                        trailing = run.text[len(stripped):]
                        run.text = empresa_nome + trailing
                        break

        for para in paras_remover:
            _remove_para(para)

        # Salva .docx em diretório temporário e converte para PDF via LibreOffice
        tmp_dir = tempfile.mkdtemp()
        try:
            docx_path = os.path.join(tmp_dir, 'termo.docx')
            doc.save(docx_path)

            subprocess.run(
                [
                    'libreoffice', '--headless',
                    f'-env:UserInstallation=file://{tmp_dir}/lo_profile',
                    '--convert-to', 'pdf',
                    '--outdir', tmp_dir,
                    docx_path,
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )

            pdf_path = os.path.join(tmp_dir, 'termo.pdf')
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        return io.BytesIO(pdf_bytes)

    @staticmethod
    def montar_descricao_de_orcamento(db: Session, orcamento_id: int) -> str:
        """Formata '3x PRODUTO_X · 1x PRODUTO_Y' — apenas itens do tipo PRODUTO."""
        from app.repositories.orcamento_repository import OrcamentoRepository
        from app.models.orcamento_model import TipoItemOrcamento
        orc = OrcamentoRepository.get_by_id(db, orcamento_id)
        if not orc or not orc.itens:
            return ""
        partes = []
        for item in orc.itens:
            if item.tipo != TipoItemOrcamento.PRODUTO:
                continue
            nome = item.nome or f"Item #{item.id}"
            qtd = item.quantidade or 1
            qtd_fmt = int(qtd) if qtd == int(qtd) else qtd
            partes.append(f"{qtd_fmt}x {nome}")
        return " · ".join(partes)
