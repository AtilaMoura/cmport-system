"""
Script de teste local do termo de garantia — sem banco de dados.
Executa com o venv do backend:
    backend\\venv\\Scripts\\activate
    python test_termo.py
"""
import os
import re
from datetime import date, datetime
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

TEMPLATE = os.path.join("backend", "app", "assets", "termo_garantia_template.docx")
OUTPUT_DOCX = "termo_teste_novo.docx"
OUTPUT_PDF = "termo_teste_novo.pdf"

# ── funções copiadas do service (independentes de DB) ─────────────────────

_EXTENSO = {3: "três", 6: "seis", 12: "doze", 24: "vinte e quatro"}

def _fmt_date(d):
    return d.strftime('%d/%m/%Y') if d else ""

def _fmt_data_extenso(d):
    meses = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    return f"{d.day} de {meses[d.month-1]} de {d.year}"

def _fmt_numero_nota(numero_raw, tipo_str):
    clean = str(numero_raw).strip().replace('.', '')
    m = re.match(r'^(\d+)', clean)
    if not m:
        return numero_raw
    num = int(m.group(1))
    digits = f"{num:09d}"
    formatted = f"{digits[0:3]}.{digits[3:6]}.{digits[6:9]}"
    sufixo = '-A' if tipo_str == 'assistencia' else '-M'
    return f"{formatted}{sufixo}"

def _para_text(para):
    return ''.join(r.text for r in para.runs)

def _merge_all(para, new_text):
    if not para.runs:
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''

def _remove_para(para):
    p = para._element
    p.getparent().remove(p)

def _remover_quebras_pagina(doc):
    from docx.oxml.ns import qn
    for para in doc.paragraphs:
        ppr = para._element.find(qn('w:pPr'))
        if ppr is not None:
            pb = ppr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                ppr.remove(pb)
        for run in para.runs:
            for br in run._r.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    run._r.remove(br)

def _remover_paragrafos_vazios_iniciais(doc):
    para_remover = []
    for para in doc.paragraphs:
        if not _para_text(para).strip():
            para_remover.append(para)
        else:
            break
    for para in para_remover:
        _remove_para(para)

def _ajustar_para_uma_pagina(doc):
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(1.5)
        section.footer_distance = Cm(0.75)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size > Pt(10):
                run.font.size = run.font.size - Pt(1)

def _set_cliente_endereco(para, nome, endereco):
    from docx.oxml.ns import qn
    runs = para.runs
    br_idx = next((i for i, r in enumerate(runs) if r._r.find(qn('w:br')) is not None), -1)
    if br_idx == -1:
        _merge_all(para, f"Cliente: {nome}. Endereço: {endereco}.")
        return
    before = runs[:br_idx]
    if len(before) >= 2:
        before[0].text = 'Cliente: '
        before[1].text = nome
        before[1].bold = True
        for r in before[2:]:
            r.text = ''
    elif len(before) == 1:
        before[0].text = f'Cliente: {nome}'
    after = runs[br_idx + 1:]
    if len(after) >= 2:
        after[0].text = 'Endereço: '
        after[1].text = endereco
        after[1].bold = True
        for r in after[2:]:
            r.text = ''
    elif len(after) == 1:
        after[0].text = f'Endereço: {endereco}'

def _set_normal_runs(para, new_value):
    normal_runs = [r for r in para.runs if r.bold is not True]
    if not normal_runs:
        return
    normal_runs[0].text = f' {new_value}'
    for r in normal_runs[1:]:
        r.text = ''

def _set_prazo(para, prazo_fmt, data_inicio, data_fim):
    runs = para.runs
    in_label = True
    intro_passed = False
    bold_prazo = []
    normal_datas = None
    for run in runs:
        if in_label:
            if run.bold is True:
                continue
            else:
                in_label = False
                intro_passed = True
                continue
        elif intro_passed and run.bold is True:
            bold_prazo.append(run)
        elif intro_passed and run.bold is not True:
            normal_datas = run
    if bold_prazo:
        bold_prazo[0].text = f'{prazo_fmt} meses'
        for r in bold_prazo[1:]:
            r.text = ''
    if normal_datas:
        normal_datas.text = f', com início em {data_inicio} e término em {data_fim}.'

# ── dados de teste (mesmo do termo 8) ─────────────────────────────────────

NOME_COND    = "Condominio Edificio Jussara"
ENDERECO     = "R. Mesquita, 117, Vila Deodoro, São Paulo"
TIPO_SERVICO = "assistencia"
NUMERO_NOTA  = "93"
NUMERO_OS    = "72696966"
PRAZO_MESES  = 12
DATA_SERVICO = date(2026, 4, 17)
DATA_INICIO  = date(2026, 4, 17)
DATA_FIM     = date(2027, 4, 17)
PRODUTO_DESC = "1x REFLETOR LED 100W SLIM EMPALUX"
EMPRESA_NOME = "CMPORT Sistemas Eletrônicos de Segurança"

# ── geração ───────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)
_remover_quebras_pagina(doc)
_remover_paragrafos_vazios_iniciais(doc)
_ajustar_para_uma_pagina(doc)

numero_nota_fmt = _fmt_numero_nota(NUMERO_NOTA, TIPO_SERVICO)
prazo_extenso   = _EXTENSO.get(PRAZO_MESES, str(PRAZO_MESES))
prazo_fmt       = f"{PRAZO_MESES:02d} ({prazo_extenso})"
hoje_str        = _fmt_data_extenso(datetime.now())
data_servico_str = _fmt_date(DATA_SERVICO)
data_inicio_str  = _fmt_date(DATA_INICIO)
data_fim_str     = _fmt_date(DATA_FIM)

paras_remover = []
for para in doc.paragraphs:
    text = _para_text(para)
    if not text.strip():
        continue

    if "Paulo," in text and len(text) < 60:
        _merge_all(para, f"São Paulo, {hoje_str}")

    elif "Cliente:" in text and "Endere" in text:
        _set_cliente_endereco(para, NOME_COND, ENDERECO)

    elif "consistente na" in text:
        for run in para.runs:
            if run.bold is True:
                run.text = PRODUTO_DESC
                run.font.size = Pt(9)
                break

    elif "execu" in text and "servi" in text and "/" in text:
        _set_normal_runs(para, data_servico_str)

    elif "Nota Fiscal:" in text:
        if numero_nota_fmt:
            _set_normal_runs(para, f"nº {numero_nota_fmt}")
        else:
            paras_remover.append(para)

    elif "Ordem de Servi" in text:
        if NUMERO_OS:
            _set_normal_runs(para, f"nº {NUMERO_OS}")
        else:
            paras_remover.append(para)

    elif "Prazo de Garantia:" in text and "garantia concedida" in text:
        _set_prazo(para, prazo_fmt, data_inicio_str, data_fim_str)

    elif "André Moreira Rosa" in text:
        for run in para.runs:
            stripped = run.text.rstrip()
            if stripped and "André" not in stripped and "Diretor" not in stripped:
                trailing = run.text[len(stripped):]
                run.text = EMPRESA_NOME + trailing
                break

for para in paras_remover:
    _remove_para(para)

doc.save(OUTPUT_DOCX)
print(f"Salvo DOCX em: {os.path.abspath(OUTPUT_DOCX)}")

print("Convertendo DOCX para PDF usando o próprio motor do Microsoft Word (docx2pdf)...")
convert(OUTPUT_DOCX, OUTPUT_PDF)
print(f"Salvo PDF em: {os.path.abspath(OUTPUT_PDF)}")
print("Pode conferir o PDF. O layout agora é 100% idêntico ao do Word!")
